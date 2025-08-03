import os
import re
import json
from datetime import datetime
import PyPDF2 as pdf
from docx import Document
import google.generativeai as genai
from google.generativeai import types
from pydantic import BaseModel
from utils import optimize_keywords, enforce_page_limit
from dotenv import load_dotenv
from streamlit import session_state as st_session
import openai

openai.api_key = os.getenv("OPENAI_API_KEY")


# Initialize Gemini client
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
model = genai.GenerativeModel("gemini-2.5-flash")

class CVOptimization(BaseModel):
    """CV optimization response model"""
    ats_score: int
    missing_keywords: list
    optimized_content: str
    suggestions: list

def extract_resume_text(uploaded_file):
    """Extract text from uploaded resume file"""
    if uploaded_file.name.endswith(".pdf"):
        reader = pdf.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    elif uploaded_file.name.endswith(".docx"):
        doc = Document(uploaded_file)
        return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    else:
        return ""

def generate_cv(resume_text, job_description, target_match, template, sections, quantitative_focus, action_verb_intensity, keyword_matching):
    """Generate optimized CV using Gemini AI"""
    
    # Build sections string
    sections_list = [section for section, include in sections.items() if include]
    sections_string = ", ".join(sections_list)
    
    # Adjust prompt based on settings
    intensity_mapping = {
        "Moderate": "moderate use of action verbs",
        "High": "strong emphasis on action verbs",
        "Very High": "maximum use of powerful action verbs"
    }
    
    matching_mapping = {
        "Conservative": "maintain authenticity while incorporating key terms",
        "Balanced": "strategically integrate job description keywords",
        "Aggressive": "maximize keyword density and exact phrase matching"
    }
    
    # Direct prompt for CV output only
    prompt = f"""
    You are a professional resume writer and an expert in ATS optimization and role alignment.

    Your job is to:
    1. Parse the candidate's resume and extract **real experience**.
    2. Analyze the job description to extract **critical keywords, tools, titles, skills, certifications, and action verbs**.
    3. Identify mismatches between the resume and JD (especially job titles like "Data Analyst" vs. "Data Engineer").
    4. Reframe the resume to match the **job role in the JD**, especially:
    - Rewrite bullet points to highlight experience adjust Real experience with the JD's Skills.
    - Emphasize **tools, platforms, pipelines, databases, programming, and architecture** relevant to the target role.
    - Add **measurable outcomes and business impact** wherever possible.
    EXECUTE UNIVERSAL CV GENERATION: Analyze JD, extract 45 ATS skills, generate 100-word summary, create 22 JD-aligned roles across all companies and ensure the entire content fits within 2 A4 pages. Use only exact wording from the JD. No paraphrasing. No personal data. Avoid repetition. Ensure perfect ATS compatibility, and quantifiable outcomes in 50%+ of roles.
    Steps:
    Extract 45 unique ATS-compliant skills from the JD using exact wording. Limit each skill to 1-2 words. Categorize into: 15 Technical Skills, 15 Soft Skills, 15 Job-Specific Competencies.
    Write a 100-word summary starting with “Applying for [exact job title]”. Include [X]+ years experience, 15+ ATS keywords, quantifiable outcomes, global exposure, and action verbs. No synonyms.
    Create 22 roles across all companies in resume: . Each role must be 10-14 words, use 1-2 ATS skills, end with a full stop. 50%+ roles must have quantifiable metrics. Avoid repeating skills across roles. Ensure progressive responsibility.


    Fabricate work experience to better align with the JD with quantitive ATS keywords.
    Do not change the job title
    Just adjust the ATS keywords, skills inside the roles and responsibility.

    Your goal is to improve this resume to achieve a **{target_match}% ATS match** with the JD.

    Generate the resume in this exact plain text format with these headers (Headers in Bold), make sure name and details are in centre:

    NAME
    Phone No | Email | Address
    # Make sure NAME and contact details are at the top, centered, and not under any section

    PROFESSIONAL SUMMARY:
    

    KEY SKILLS:
    Skill 1, Skill 2.....

    WORK EXPERIENCE:(keep the dates in the same format as given in resume)
    Company | Role | Dates
    • Bullet 1
    • Bullet 2

    EDUCATION:
    • Degree | Institution | Year(keep the dates in the same format as given in resume)

    PROJECTS:(if any)
    Project Name 1
    • Bullet 1
    • Bullet 2
    
    Project Name 2
    • Bullet 1
    • Bullet 2

    CERTIFICATIONS:(If any)

    Resume Content:
    {resume_text}

    Job Description:
    {job_description}
    """

    prompt_2 = f"""
    You are an expert ATS resume writer with deep knowledge of applicant tracking systems and recruitment best practices.

    GOAL:
    Rewrite the candidate's resume to achieve {target_match if 'target_match' in locals() else '90'}% ATS match with the job description while maintaining authenticity, professional quality, and ATS compliance. Make sure roles should not be more than 22.

    INPUT VALIDATION:
    - If resume_text is missing key sections (contact, experience), note what's missing and proceed with available information
    - If job_description is incomplete or too brief, focus on industry-standard keywords for the role

    CRITICAL ATS TECHNICAL REQUIREMENTS:
    ✔ Use standard fonts only: Calibri (11-12pt)
    ✔ Set margins: 0.5-1 inch on all sides
    ✔ Use single line spacing with 6pt spacing after paragraphs
    ✔ Avoid: Tables, text boxes, graphics, images, headers/footers, columns
    ✔ Use standard bullet points (•) consistently

    CONTENT AUTHENTICITY RULES:
    ✔ Include NAME and contact details at the top, centered
    ✔ Keep ALL companies, job titles, and dates from original resume unchanged
    ✔ Keep educational institutions and degrees exactly as provided
    ✔ Include EVERY work experience from the original resume - do not omit any roles
    ✔ Maintain chronological order of all positions (reverse chronological)
    ✔ Do NOT invent companies, roles, or qualifications
    ✔ Enhance responsibilities and achievements within existing roles only
    ✔ If employment gaps exist, do not fill them - maintain chronological accuracy

    FORMATTING STANDARDS:
    ✔ Date format: MM/YYYY consistently throughout
    ✔ Contact format: Phone | Email | City, State (optional ZIP)
    ✔ Use consistent capitalization for section headers
    ✔ Each bullet point: 10-14 words (flexibility for complex achievements)
    ✔ End every bullet point with a period
    ✔ Use reverse chronological order for experience

    CONTENT ENHANCEMENT REQUIREMENTS:
    ✔ Generate exactly 22 bullet points total across ALL work experience roles:
    - Distribute bullets proportionally based on role relevance to target position
    - Most recent/relevant role: 5 bullets
    - Second most recent/relevant: 4 bullets  
    - Remaining roles: 3, 2 or 1 bullets each in decreasing order
    - Adjust distribution based on number of roles but maintain 22 total
    ✔ Add quantifiable metrics to 60% of bullet points minimum (approximately 13-14 bullets):
    - Percentages (increased efficiency by 25%)
    - Dollar amounts ($2M budget managed)
    - Numbers (managed team of 12)
    - Timeframes (reduced processing time by 3 days)
    ✔ If original achievements lack numbers, create realistic estimates based on role level and industry standards
    ✔ Each bullet point must include 1-2 relevant JD keywords naturally integrated
    ✔ Prioritize exact keyword matches over synonyms
    ✔ Balance hard skills (70%) and soft skills (30%) in keyword integration

    PAGE LENGTH MANAGEMENT:
    ✔ Target: 2 pages maximum

    QUALITY CONTROLS:
    ✔ Ensure professional tone throughout
    ✔ Avoid keyword stuffing (max 2-3 keywords per bullet)
    ✔ Verify all company names and job titles remain unchanged
    ✔ Check for spelling and grammar errors
    ✔ Ensure logical flow and readability
    ✔ Maintain consistency in tense (past tense for previous roles, present for current)

    KEYWORD STRATEGY:
    ✔ Extract 45-50 most relevant keywords from job description
    ✔ Prioritize: Technical skills, software, certifications, industry terms
    ✔ Include variations where natural (e.g., "manage" and "management")
    ✔ Distribute keywords across all sections, not just skills
    ✔ Ensure keywords appear in context, not as isolated lists

    BULLET POINT DISTRIBUTION STRATEGY:
    ✔ Total bullet points required: EXACTLY 22 across ALL work experience roles
    ✔ CRITICAL: Include every single work experience from the original resume
    ✔ Distribution guidelines based on total number of roles in original resume:
    - 1 role: All 22 bullets for that role
    - 2 roles: 11 bullets each
    - 3 roles: 8 + 7 + 7 bullets (most recent/relevant gets 8)
    - 4 roles: 6 + 6 + 5 + 5 bullets (prioritize by relevance/recency)
    - 5 roles: 5 + 5 + 4 + 4 + 4 bullets
    - 6 roles: 4 + 4 + 4 + 3 + 3 + 4 bullets
    - 7+ roles: Distribute as 4-3-3-3-3-3-3 pattern, adjust as needed
    ✔ Prioritize bullet allocation: Most recent and JD-relevant roles get more bullets
    ✔ Every role must have minimum 3 bullets, maximum 8 bullets
    ✔ Never omit or combine roles - each original position gets its own section

    EDGE CASE HANDLING:
    - Entry-level candidates: Focus on education, projects, internships, relevant coursework
    - Career changers: Emphasize transferable skills and relevant projects
    - Employment gaps: Do not fabricate - focus on remaining experience
    - Limited quantifiable achievements: Use industry benchmarks and reasonable estimates
    - NEVER omit, combine, or merge work experiences from original resume
    - If candidate has 7+ roles, still include all but distribute bullets strategically (3-4 per role)
    - For very brief roles (under 6 months), still include with minimum 3 bullets
    - Maintain complete work history integrity while optimizing for ATS

    OUTPUT FORMAT (STRICT - DO NOT OMIT ANY SECTION):

    [CANDIDATE NAME]
    Phone | Email | City, State

    PROFESSIONAL SUMMARY:
    [100-120 word summary starting with "Accomplished [job title/field] professional applying for [exact job title from JD]..." Include 3-4 key JD keywords naturally.]

    KEY SKILLS:
    [List 45-50 ATS keywords from JD, comma-separated, prioritizing exact matches. Mix technical skills, software, methodologies, and relevant soft skills.]

    WORK EXPERIENCE:
    [Include ALL work experiences from original resume in reverse chronological order]

    [Company Name] | [Job Title] | [MM/YYYY - MM/YYYY or Present]
    • [Achievement with metric and 1-2 JD keywords].
    • [Responsibility with quantified outcome and relevant keywords].
    • [Impact statement with numbers and key terms from JD].
    [Repeat format for EVERY role from original resume - distribute 22 total bullets across all positions]

    [Company Name] | [Job Title] | [MM/YYYY - MM/YYYY]
    • [Achievement with metric and 1-2 JD keywords].
    • [Responsibility with quantified outcome and relevant keywords].
    • [Impact statement with numbers and key terms from JD].
    [Continue for each role - ensure no original work experience is omitted]

    EDUCATION:
    [Degree] | [Institution] | [Year]
    [Additional certifications, relevant coursework, or academic achievements if space permits]

    PROJECTS: (if applicable)
    [Project Name] | [Technologies/Skills Used] | [MM/YYYY]
    • [Brief description with outcomes and relevant keywords].

    CERTIFICATIONS: (if any)
    [Certification Name] | [Issuing Organization] | [Year/Expiration]

    
    ERROR HANDLING:
    - If resume sections are missing, create reasonable structure with available information
    - If job description lacks detail, use industry-standard keywords for the role type
    - If unable to reach target match percentage, aim for highest possible match while maintaining authenticity

    POST-PROCESSING CHECKLIST:
    □ All original company names and job titles preserved
    □ Contact information complete and properly formatted
    □ EXACTLY 22 bullet points in work experience section
    □ Keywords naturally integrated, not stuffed
    □ Quantified achievements in 60%+ of bullets (13-14 out of 22)
    □ Professional tone maintained throughout
    □ Spelling and grammar checked
    □ Page count within limits
    □ ATS-friendly formatting applied
    □ File should be saved as .docx format

    INPUT DATA:
    Resume Text:
    {resume_text if 'resume_text' in locals() else '[RESUME TEXT MISSING - Please provide resume content]'}

    Job Description:
    {job_description if 'job_description' in locals() else '[JOB DESCRIPTION MISSING - Please provide job posting details]'}

    Target Match Percentage: {target_match if 'target_match' in locals() else '85'}%

    FINAL INSTRUCTION: Return ONLY the formatted resume following the exact structure above. Do not include explanations, meta-commentary, or additional text. End with: "NOTE: Save this resume as a .docx file for optimal ATS compatibility."
    """

    
    try:
        # ✅ OpenAI Flow
        if st_session.get("ai_model") == "openai":
            response = openai.chat.completions.create(
                model="gpt-4.1",
                messages=[
                    {"role": "system", "content": "You are a professional resume writer."},
                    {"role": "user", "content": prompt_2}
                ],
                temperature=0.2
            )
            optimized_cv = response.choices[0].message.content
            # ✅ Post-process and return
            optimized_cv = clean_cv_content(optimized_cv)
            optimized_cv = enforce_page_limit(optimized_cv)
            return optimized_cv.strip()

        # ✅ Gemini Flow
        else:

            response = model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.2,
                "response_mime_type": "text/plain"
            }
            )
        
        # Handle different response conditions
        if not response:
            raise Exception("No response received from AI")
        
        if response.candidates and len(response.candidates) > 0:
            candidate = response.candidates[0]
            if candidate.finish_reason.name == 'MAX_TOKENS':
                # Try to get partial content
                if candidate.content and candidate.content.parts:
                    partial_text = ""
                    for part in candidate.content.parts:
                        if hasattr(part, 'text') and part.text:
                            partial_text += part.text
                    if partial_text:
                        optimized_cv = partial_text
                    else:
                        raise Exception("MAX_TOKENS reached and no partial content available")
                else:
                    raise Exception("MAX_TOKENS reached and no content available")
            elif not response.text:
                raise Exception("AI response was empty")
            else:
                optimized_cv = response.text
        else:
            raise Exception("No candidates in response")
        

        
        # Clean up the response
        optimized_cv = clean_cv_content(optimized_cv)
        optimized_cv = enforce_page_limit(optimized_cv)

        from utils import extract_keywords_from_text

        jd_keywords = extract_keywords_from_text(job_description)

        def bold_keywords_in_work_exp(cv_text, keywords):
            if "WORK EXPERIENCE:" not in cv_text:
                return cv_text

            parts = cv_text.split("WORK EXPERIENCE:")
            before = parts[0]
            after = parts[1]

            lines = after.split('\n')
            bolded_lines = []
            for line in lines:
                if line.startswith("•") or "|" in line:
                    for kw in keywords:
                        pattern = r'\b(' + re.escape(kw) + r')\b'
                        line = re.sub(pattern, r'**\1**', line, flags=re.IGNORECASE)
                bolded_lines.append(line)

            return before + "WORK EXPERIENCE:\n" + '\n'.join(bolded_lines)

        optimized_cv = bold_keywords_in_work_exp(optimized_cv, jd_keywords)

        return optimized_cv.strip()
        
    except Exception as e:
        raise Exception(f"Failed to generate CV: {str(e)}")

def generate_cover_letter(resume_text, job_description):
    """Generate cover letter using Gemini AI"""
    
    prompt = f"""
    You are an expert ATS-optimized cover letter writer.
    
    Objective:
    Generate a personalized, professional cover letter that achieves **90%+ ATS compatibility** and aligns precisely with the provided Job Description.
    
    Rules:
    - Dynamically adjust keyword placement to ensure **high ATS score**.
    - Start with: “Hello Hiring Manager,” and include the line: “I am applying for the [exact job title] position.”
    - Use a tone that reflects professionalism and enthusiasm.
    
    Structure:
    1. **Paragraph 1**: Express genuine enthusiasm using the company's mission and JD language. Include company-specific values, vision, and relevant projects to show personalization.
    2. **Paragraph 2**: Align with the top 5 responsibilities in the JD. Provide **metrics-rich accomplishments** from the resume that demonstrate capability. Integrate at least **10 relevant keywords** from the JD (e.g., Python, SQL, machine learning, A/B testing, scikit-learn, AWS, Snowflake, dashboards, Spark, data-driven decision-making).
    3. **Paragraph 3**: Highlight **2–3 JD outcome-based goals** (e.g., predictive models, actionable insights, collaboration with cross-functional teams) using similar phrasing and past success examples. Include any statistical analysis, testing, or ML exposure.
    4. **Paragraph 4**: Reaffirm **2 key JD priorities**. Close by offering measurable value in JD terms. Request an interview and include a polite sign-off.
    
    Additional Requirements:
    - Use **identical terminology from the JD** wherever possible (e.g., "predictive models," "statistical analyses," "machine learning frameworks").
    - Mention **preferred skills** if applicable (e.g., AWS, Snowflake, Power BI).
    - Keep tone formal yet engaging, max 4 paragraphs.
    - After sign-off, include candidate's email and phone number (extract from resume).
    
    Inputs:
    Resume:
    {resume_text}
    
    Job Description:
    {job_description}
    
    Output:
    Generate the final cover letter in **plain text** format without extra commentary.
    """

    try:
        if st_session.get("ai_model") == "openai":
            response = openai.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a professional cover letter writer."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            cover_letter = response.choices[0].message.content
            return re.sub(r'\*{1,2}', '', cover_letter).strip()

        else:

            response = model.generate_content(model="gemini-2.5-flash", contents=prompt)
        if not response or not response.text:
            raise Exception("AI response was empty")

        return re.sub(r'\*{1,2}', '', response.text).strip()

    except Exception as e:
        raise Exception(f"Failed to generate cover letter: {str(e)}")

def clean_cv_content(content):
    """Clean and format CV content"""
    if not content:
        return "Error: No content received from AI"
    
    # Remove markdown formatting
    content = re.sub(r'\*\*', '', content)
    content = re.sub(r'__', '', content)
    
    # Remove excessive whitespace
    content = re.sub(r'\n{3,}', '\n\n', content)
    
    # Remove any hidden markers
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)
    
    # Ensure proper section formatting
    content = re.sub(r'^([A-Z][A-Z\s]+):', r'\n\1:', content, flags=re.MULTILINE)
    
    return content.strip()

def analyze_cv_ats_score(cv_content, job_description):
    """Analyze CV ATS compatibility score using Gemini AI"""
    
    prompt = f"""
    You are an ATS analysis expert.
    
    Analyze the CV against the job description and provide:
    1. ATS compatibility score (0-100)
    2. Keyword match percentage
    3. Missing critical keywords
    4. Specific improvement suggestions
    
    Return JSON format:
    {{
        "ats_score": number,
        "keyword_match": number,
        "missing_keywords": [list],
        "suggestions": [list]
    }}
    
    CV Content:
    {cv_content}
    
    Job Description:
    {job_description}
    """
    
    try:
        
        
        response = model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.2,
                "response_mime_type": "application/json"
            }
            )

        if not response or not response.text:
            raise Exception("AI response was empty or None")
        
        try:
            parsed = json.loads(response.text)
        except Exception as parse_err:
            raise Exception(f"Invalid JSON response from Gemini: {response.text}")

        return {
            "score": parsed.get("ats_score", 0),
            "keyword_match": parsed.get("keyword_match", 0),
            "missing_keywords": parsed.get("missing_keywords", []),
            "suggestions": parsed.get("suggestions", [])
        }

    except Exception as e:
        # Final fallback if AI fails entirely
        return {
            "score": 0,
            "keyword_match": 0,
            "missing_keywords": [],
            "suggestions": [f"Error analyzing CV: {str(e)}"]
        }

def extract_key_metrics(cv_content):
    """Extract quantifiable metrics from CV"""
    # Pattern to find numbers and percentages
    metrics_pattern = r'(\d+(?:\.\d+)?(?:%|K|M|B|k|m|b|\+|,\d+)*)'
    
    metrics = re.findall(metrics_pattern, cv_content)
    
    return {
        'total_metrics': len(metrics),
        'metrics_found': metrics,
        'quantification_score': min(100, len(metrics) * 5)  # 5 points per metric, max 100
    }

def enhance_action_verbs(content, intensity="High"):
    """Enhance action verbs in CV content"""
    
    action_verbs = {
        "Moderate": [
            "managed", "developed", "created", "implemented", "led", "coordinated",
            "designed", "analyzed", "improved", "organized", "planned", "supervised"
        ],
        "High": [
            "spearheaded", "orchestrated", "revolutionized", "transformed", "pioneered",
            "architected", "optimized", "streamlined", "accelerated", "amplified"
        ],
        "Very High": [
            "catapulted", "revolutionized", "masterminded", "propelled", "dominated",
            "commanded", "conquered", "devastated", "obliterated", "annihilated"
        ]
    }
    
    # This would be implemented with more sophisticated text processing
    # For now, return the content as-is
    return content

def generate_interview_qa(resume_text, job_description):
    """Generate interview Q&A using Gemini AI"""
    prompt = f"""
    You are an expert career coach and interviewer.

    TASK:
    Generate **exactly 20 interview questions and answers** for the candidate based on their resume and the job description.

    ✅ Structure:
    - 8 Behavioral questions (fitment, company, teamwork, problem-solving, adaptability)
    - 12 Technical questions based on ATS keywords, JD tools, frameworks, and skills.

    ✅ Format STRICTLY:
    Q1: [Behavioral Question]
    A1:
    - Point 1
    - Point 2
    - Point 3
    - Point 4
    - Point 5
    - Point 6

    Q2: [Next Question]
    A2:
    - ...

    ✅ Rules:
    - All answers MUST have **6 bullet points minimum**.
    - No repetition of questions or answers.
    - Technical questions should be advanced and role-specific.
    - Cover JD-specific tools, methods, and problem scenarios.
    - Include the most important ATS keywords in both questions and answers.

    Resume:
    {resume_text}

    Job Description:
    {job_description}
    """
    try:
        if st_session.get("ai_model") == "openai":
            response = openai.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert career coach and interviewer."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3
            )
            return response.choices[0].message.content

        else:

            response = model.generate_content(model="gemini-2.5-flash", contents=prompt)
        if not response or not response.text:
            raise Exception("AI response was empty")

        return response.text

    except Exception as e:
        raise Exception(f"Failed to generate Q&A: {str(e)}")


def export_interview_qa(content):
    """Export Q&A content as PDF and DOCX"""
    from io import BytesIO
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from docx import Document

    # PDF Export
    pdf_buffer = BytesIO()
    doc = SimpleDocTemplate(pdf_buffer)
    styles = getSampleStyleSheet()
    story = []
    for line in content.split('\n'):
        if line.strip():
            story.append(Paragraph(line.strip(), styles['Normal']))
            story.append(Spacer(1, 12))
    doc.build(story)
    pdf_buffer.seek(0)

    # DOCX Export
    docx_buffer = BytesIO()
    word_doc = Document()
    for line in content.split('\n'):
        if line.strip():
            word_doc.add_paragraph(line.strip())
    word_doc.save(docx_buffer)
    docx_buffer.seek(0)

    return pdf_buffer, docx_buffer
